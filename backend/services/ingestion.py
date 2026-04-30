"""Ingestion pipeline: parse, chunk, embed, and persist documents to ChromaDB."""

from __future__ import annotations

import logging
import uuid
from dataclasses import dataclass, field
from pathlib import Path

import chromadb
import pymupdf
import tiktoken
from chromadb.api.models.Collection import Collection
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer

from backend.config import get_settings

logger = logging.getLogger(__name__)

COLLECTION_NAME = "documents"
_EMBEDDER: SentenceTransformer | None = None
_CHROMA_CLIENT: chromadb.api.ClientAPI | None = None
_COLLECTION: Collection | None = None

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
_ENCODING = tiktoken.get_encoding("cl100k_base")


@dataclass(frozen = True)
class ParsedPage:
    """A single page of parsed text."""

    text: str
    page_number: int


@dataclass(frozen = True)
class DocumentChunk:
    """A token-bounded chunk ready for embedding."""

    text: str
    doc_id: str
    filename: str
    page_number: int
    chunk_index: int
    token_count: int

    def metadata(self) -> dict[str, str | int]:
        """ChromaDB-compatible metadata dict (scalar values only)."""
        return {"doc_id": self.doc_id, "filename": self.filename, "page_number": self.page_number,
            "chunk_index": self.chunk_index, "token_count": self.token_count, }


@dataclass(frozen = True)
class ParsedDocument:
    """A parsed document with identity and page list."""

    doc_id: str
    filename: str
    pages: list[ParsedPage] = field(default_factory = list)


def parse_document(path: str | Path, doc_id: str | None = None) -> ParsedDocument:
    """Parse a PDF, DOCX, or TXT file into a ParsedDocument.

    Raises:
        FileNotFoundError: path does not exist.
        ValueError: extension is not supported.
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"No such file: {p}")

    ext = p.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported extension {ext!r}; expected one of {sorted(SUPPORTED_EXTENSIONS)}")

    if ext == ".pdf":
        pages = _parse_pdf(p)
    elif ext == ".docx":
        pages = _parse_docx(p)
    else:
        pages = _parse_txt(p)

    return ParsedDocument(doc_id = doc_id or uuid.uuid4().hex, filename = p.name, pages = pages, )


def _parse_pdf(path: Path) -> list[ParsedPage]:
    pages: list[ParsedPage] = []
    with pymupdf.open(path) as doc:
        for i, page in enumerate(doc, start = 1):
            text = page.get_text("text").strip()
            if text:
                pages.append(ParsedPage(text = text, page_number = i))
    return pages


def _parse_docx(path: Path) -> list[ParsedPage]:
    # python-docx does not expose true page breaks from the XML, so the whole
    # document is surfaced as page 1. Upstream consumers should not depend on
    # page-level granularity for DOCX inputs.
    doc = DocxDocument(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [ParsedPage(text = text, page_number = 1)] if text else []


def _parse_txt(path: Path) -> list[ParsedPage]:
    text = path.read_text(encoding = "utf-8", errors = "replace").strip()
    return [ParsedPage(text = text, page_number = 1)] if text else []


def chunk_document(document: ParsedDocument, chunk_size: int | None = None, chunk_overlap: int | None = None, ) -> list[
    DocumentChunk]:
    """Split a parsed document into token-bounded chunks with overlap.

    Chunks never cross page boundaries, so each chunk carries the page number
    it came from. Sizes default to the values in `backend.config.Settings`.
    """
    settings = get_settings()
    size = chunk_size if chunk_size is not None else settings.chunk_size
    overlap = chunk_overlap if chunk_overlap is not None else settings.chunk_overlap

    if size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap < 0 or overlap >= size:
        raise ValueError("chunk_overlap must satisfy 0 <= overlap < chunk_size")

    step = size - overlap
    chunks: list[DocumentChunk] = []
    global_index = 0

    for page in document.pages:
        tokens = _ENCODING.encode(page.text)
        if not tokens:
            continue
        for start in range(0, len(tokens), step):
            window = tokens[start: start + size]
            if not window:
                break
            chunks.append(
                DocumentChunk(text = _ENCODING.decode(window), doc_id = document.doc_id, filename = document.filename,
                    page_number = page.page_number, chunk_index = global_index, token_count = len(window), ))
            global_index += 1
            if start + size >= len(tokens):
                break

    logger.info("Chunked %s: %d pages -> %d chunks (size=%d overlap=%d)", document.filename, len(document.pages),
        len(chunks), size, overlap, )
    return chunks


def ingest_file(path: str | Path) -> list[DocumentChunk]:
    """Parse + chunk a file in one call (no persistence)."""
    return chunk_document(parse_document(path))


@dataclass(frozen = True)
class IngestResult:
    """Outcome of a full parse -> chunk -> embed -> persist cycle."""

    doc_id: str
    filename: str
    num_pages: int
    num_chunks: int


def get_embedder() -> SentenceTransformer:
    """Lazily load and cache the sentence-transformers model."""
    global _EMBEDDER
    if _EMBEDDER is None:
        name = get_settings().embedding_model
        logger.info("Loading embedding model: %s", name)
        _EMBEDDER = SentenceTransformer(name)
    return _EMBEDDER


def embed_texts(texts: list[str]) -> list[list[float]]:
    """Embed a batch of texts. Returns plain Python lists for Chroma."""
    if not texts:
        return []
    vectors = get_embedder().encode(texts, convert_to_numpy = True, show_progress_bar = False,
        normalize_embeddings = True, )
    return vectors.tolist()


def get_collection() -> Collection:
    """Return the shared ChromaDB collection, creating it if needed."""
    global _CHROMA_CLIENT, _COLLECTION
    if _COLLECTION is None:
        settings = get_settings()
        Path(settings.chroma_persist_dir).mkdir(parents = True, exist_ok = True)
        _CHROMA_CLIENT = chromadb.PersistentClient(path = settings.chroma_persist_dir)
        _COLLECTION = _CHROMA_CLIENT.get_or_create_collection(name = COLLECTION_NAME,
            metadata = {"hnsw:space": "cosine"}, )
    return _COLLECTION


def store_chunks(chunks: list[DocumentChunk]) -> int:
    """Embed and persist chunks. Returns the number stored."""
    if not chunks:
        return 0

    ids = [f"{c.doc_id}:{c.chunk_index}" for c in chunks]
    documents = [c.text for c in chunks]
    metadatas = [c.metadata() for c in chunks]
    embeddings = embed_texts(documents)

    get_collection().upsert(ids = ids, embeddings = embeddings, documents = documents, metadatas = metadatas, )
    logger.info("Stored %d chunks for %s", len(chunks), chunks[0].filename)
    # Lazy import — bm25 imports back into ingestion for the collection accessor.
    from backend.services import bm25 as _bm25
    _bm25.mark_dirty()
    return len(chunks)


def ingest_document(path: str | Path, doc_id: str | None = None) -> IngestResult:
    """Full pipeline: parse -> chunk -> embed -> persist."""
    parsed = parse_document(path, doc_id = doc_id)
    chunks = chunk_document(parsed)
    stored = store_chunks(chunks)
    return IngestResult(doc_id = parsed.doc_id, filename = parsed.filename, num_pages = len(parsed.pages),
        num_chunks = stored, )


@dataclass(frozen=True)
class DocumentSummary:
    """Aggregated view of a single document in the corpus."""

    doc_id: str
    filename: str
    num_chunks: int
    num_pages: int


def list_documents() -> list[DocumentSummary]:
    """Return one summary per ingested document, sorted by filename."""
    collection = get_collection()
    if collection.count() == 0:
        return []

    rows = collection.get(include=["metadatas"])
    by_doc: dict[str, dict] = {}
    for meta in rows["metadatas"] or []:
        doc_id = str(meta.get("doc_id", ""))
        if not doc_id:
            continue
        entry = by_doc.setdefault(
            doc_id,
            {"filename": str(meta.get("filename", "")), "chunks": 0, "max_page": 0},
        )
        entry["chunks"] += 1
        entry["max_page"] = max(entry["max_page"], int(meta.get("page_number", 0)))

    summaries = [
        DocumentSummary(
            doc_id=doc_id,
            filename=v["filename"],
            num_chunks=v["chunks"],
            num_pages=v["max_page"],
        )
        for doc_id, v in by_doc.items()
    ]
    summaries.sort(key=lambda s: (s.filename, s.doc_id))
    return summaries


def delete_document(doc_id: str) -> int:
    """Delete every chunk belonging to `doc_id`. Returns the number deleted."""
    collection = get_collection()
    existing = collection.get(where={"doc_id": doc_id}, include=[])
    n = len(existing.get("ids", []) or [])
    if n == 0:
        return 0
    collection.delete(where={"doc_id": doc_id})
    logger.info("Deleted %d chunks for doc_id=%s", n, doc_id)
    from backend.services import bm25 as _bm25
    _bm25.mark_dirty()
    return n
